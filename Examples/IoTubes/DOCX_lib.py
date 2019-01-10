
from docx import Document
from docx.shared import Inches
from graph_lib import gl
import pandas as pd
import datetime as dt
import os
# Creare new Word Document

class CleaningReport():
    def __init__(self):
        ## Hardcoded param:
        self.IoTubes_logo_path ="./IoTubes/images_IoTubes/IoTubes_logo.png"
        self.sensor_images_path = None #"./sensor_images/"
        self.report_path = None 
        ## Main config options
        self.report_type = "Cleaning Report";

        ## Date of the report generated
        self.date_report_generation = None 
        # Date when the cleaning process took place
        self.date_cleaning = dt.datetime.now().strftime("%B %d, %Y")

        # Client company
        self.client_ID = "CJK Steel"
        # ID of the machine that performed the cleanning
        self.machine_ID = None
        # ID of the cleaning process
        self.cleaning_ID = None
        # ID of the piping system to be cleaned
        self.piping_ID = None
        
        # Cleaning Responsibles. List of cleaning responsibles.
        self.Responsibles =  None #["Carsten", "Ole P. Ness"];
        # FDA Cleaning requirements 
        self.FDA_cleaning_procedure = None # "FDA53531984EX"
    
        ## CLEANING PROCESS DATA
        self.sensors_data_pd = None

        ## Document variables
        self.document = None
        

    def load_administration_data(self, administration_data = [None, None, None, None,None]):
        # Loading the Administration Data
        # This should feel the information for the company
        # With the cleaning ID we should be able to get all info
        # hard code in worst scenario.
        self.machine_ID, self.piping_ID, self.cleaning_ID,self.FDA_cleaning_procedure,self.Responsibles  = administration_data;
        
        pass

    def load_sensors_data(self, data, column_names = ["Temp","PH","Pressure","Conductivity"]):
        # Loading the cleaning process Data.
        # TODO: Working under the assumption that all the sensors are sampled at the same time.
        # Otherwise we would have different time labels and the SQL dabase would need a table per sensor to be efficient,.
        # data is the pandas dataframe table. 
        # The data should be provided externally, probably with a DDBB call.
        self.time_name = "TS"  # The name of the time 
        data_timestamps, data_sensors  = data
        # column_names = ["Temp","PH","Pressure","Conductivity"] #"TS",
        
        df = pd.DataFrame(data_sensors, columns = column_names)
        df[self.time_name] = data_timestamps
    
        self.sensors_data_pd = df
        self.sensor_names = column_names;

        
        pass
    
    def process_cleaning_data(self, Monitors):
        # This function should tell us if there has been any violation in the data.
        # For this purpose, it needs the Data and the Monitoring information
        # Monitors is a list of Monitor objects, the "ID" of the object would the columns of data
        self.Monitors = Monitors;
        # TODO

    def generate_images(self, folder_path):
        self.sensor_images_path = folder_path
        # This function is suposed to generate the cleaning images.
        # Basically the time series with the sensors. 
        # It uses the pandas dataframe self.sensors_data_pd
        # The images are stored in self.sensor_images_path using as names the column names of the data
        for sensor_column in self.sensor_names:
            ## TODO: Big task... manage different windows in gl library
#            gl.init_figure()
#            gl.plot(self.sensors_data_pd[self.time_name],self.sensors_data_pd[sensor_column])
            path_image = self.sensor_images_path + sensor_column + ".png"
            gl.savefig( path_image,
               dpi = 100, sizeInches = [])  # [2*8, 2*3]
            
            
        pass


    def add_FDA_requirements_table(self):
        # Just created the word table showing the parameters of the cleaning process
        # TODO: I need to add it somethow.
        
        table = self.document .add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Units'
        hdr_cells[2].text = 'Desired value'
        hdr_cells[3].text = 'Warning Range'
        hdr_cells[4].text = 'Error Range'
        
#        for item in recordset:
        for monitor in self.Monitors:
            row_cells = table.add_row().cells
            row_cells[0].text = monitor.name
            row_cells[1].text = monitor.units
            row_cells[2].text = str(monitor.desired_value)
            row_cells[3].text =  str( monitor.range_warning)
            row_cells[4].text = str(monitor.range_stop)

        
    def create_document (self, report_path = 'presentation.docx'):
        # This function will create the document, assuming all the data is loaded

        # Create word document
        document = Document()
        self.document = document
        ######################## FIRST PAGE #####################
        # Add IoTubes logo
        document.add_picture(self.IoTubes_logo_path, width=Inches(3.0))

        # Create the heading 
        document.add_heading('Cleaning Report IoTubes', 0)

        # Create Initial text describing the document
        p = document.add_paragraph()

        # Now we add different texts of the first paragraph.
        initial_text = 'This document contains the automatically generated report regarding the cleaning of the piping system '
        p.add_run(initial_text)
        initial_text = self.piping_ID + " "
        p.add_run(initial_text).bold = True

        initial_text = "carried out by the cleaning machine "
        p.add_run(initial_text)
        initial_text = self.machine_ID + " "
        p.add_run(initial_text).bold = True

        initial_text = "the day "
        p.add_run(initial_text)
        initial_text = str(self.date_cleaning ) + " "
        p.add_run(initial_text).bold = True

        initial_text = "according to the FDA standard procedure "
        p.add_run(initial_text)
        initial_text = str(self.FDA_cleaning_procedure) + ". "
        p.add_run(initial_text).bold = True

        initial_text = "The process was carried out by: "
#        for res in self.Responsibles:
#            initial_text += res " ";
        initial_text += self.Responsibles[0] + " and " + self.Responsibles[1]
        initial_text = initial_text + "."
        p.add_run(initial_text)
        
        ############ Cleaning Requirements #########
        ## State the required time, pressure, temperature, PH... and their warning and errors sent.
        ## In a Table
        document.add_heading('Cleaning Requirements', level=1)
        text = 'The requirements for the FDA cleaning procedure %s are exposed in the following table:' % str(self.FDA_cleaning_procedure)
        p = document.add_paragraph(text)
        
        self.add_FDA_requirements_table()
        
#        document.add_paragraph(
#            'first item in unordered list', style='ListBullet'
#        )
#        document.add_paragraph(
#            'first item in ordered list', style='ListNumber'
#        )

        ############ Summary of the cleaning #########
        ## State how the process went in general, like
        """
        Duration, did it meet the FDA requirements ? Warnings issued ? 
        """

        ############ SECOND PAGE: Cleaning Process Report #########
        document.add_page_break()
        document.add_heading('Cleaning Report', level=1)

        for sensor_name in self.sensor_names:
            document.add_heading(sensor_name + ' report', level=2)
            text = "The next Figure shows the levels of "+sensor_name + " during the cleaning procedure. "
            p = document.add_paragraph(text)
            
            incident_PF = False
            if (incident_PF):
                p.run("There was a warning due to the violation of boundaries")

            # Add PH image
            path_image = self.sensor_images_path + sensor_name + ".png"
            document.add_picture(path_image, width=Inches(5.74))
            
        ### Save document in the end 
        
        doc_path =  report_path # 'presentation.docx'
        document.save(doc_path)
        os.system("chmod 777 " + doc_path)

